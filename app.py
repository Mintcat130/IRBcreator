import streamlit as st
import anthropic
import io
import requests
from scholarly import scholarly
import json
import re
import uuid
import streamlit.components.v1 as components
from collections import defaultdict
from docx import Document
from io import BytesIO
from difflib import SequenceMatcher
from pdfminer.high_level import extract_text

#연구계획서 ID 생성
def generate_research_id():
    return str(uuid.uuid4())

#세션 상태 초기화
def reset_session_state():
    keys_to_keep = ['api_key', 'anthropic_client']
    for key in list(st.session_state.keys()):
        if key not in keys_to_keep:
            del st.session_state[key]
    st.session_state.current_research_id = generate_research_id()
    st.session_state.section_contents = {}

if 'show_full_content' not in st.session_state:
    st.session_state.show_full_content = False

# 섹션 내용 저장
def save_section_content(section, content):
    if 'research_data' not in st.session_state:
        st.session_state.research_data = {}
    if st.session_state.current_research_id not in st.session_state.research_data:
        st.session_state.research_data[st.session_state.current_research_id] = {}
    st.session_state.research_data[st.session_state.current_research_id][section] = content

# 섹션 내용 로드
def load_section_content(section):
    try:
        if 'research_data' in st.session_state and st.session_state.current_research_id in st.session_state.research_data:
            return st.session_state.research_data[st.session_state.current_research_id].get(section, "")
    except Exception as e:
        st.error(f"섹션 내용을 불러오는 중 오류가 발생했습니다: {e}")
    return ""


# 페이지 설정을 코드 최상단에 추가
st.set_page_config(page_title="📖IRB 연구계획서 도우미", page_icon="📖")

# 시스템 프롬프트
SYSTEM_PROMPT = """
당신은 병리학 분야의 연구 전문가로서 행동하는 AI 조수입니다. 당신의 역할은 사용자가 연구계획서를 작성하는 데 도움을 주는 것입니다. 사용자는 연구계획서의 특정 항목에 대한 정보를 제공할 것이며, 당신은 이를 바탕으로 해당 항목을 작성해야 합니다.

사용자가 제공한 정보를 주의 깊게 분석하고, 당신의 병리학 연구에 대한 전문 지식을 활용하여 요청된 연구계획서 섹션을 작성하세요. 다음 지침을 따르세요:

1. 사용자가 제공한 정보를 최대한 활용하세요.
2. 필요한 경우, 병리학 연구에 대한 당신의 지식을 사용하여 정보를 보완하세요.
3. 연구계획서 섹션의 구조와 형식을 적절히 유지하세요.
4. 명확하고 전문적인 언어를 사용하세요.
5. 필요한 경우 적절한 참고문헌이나 인용을 포함하세요.
6. 환자의 권익 보호를 최우선으로 고려하여 작성하세요.
7. 자료 보안 및 정보 보호에 대한 내용을 적절히 포함시키세요.
8. 연구방법 및 내용을 너무 구체적으로 작성하지 마세요. 연구 진행 중 방법이나 통계 기법의 변경이 가능하도록 유연성을 유지하세요.

한국어로 작성하되 의학 용어나 통계용어는 괄호 안에 영어 원문을 포함시키세요. 한국어로 번역이 불가능한 고유명사는 영어 그대로 적으세요. 예를 들어, "엽상종양(Phyllodes tumor)", "Student T-검정(Student T-test)"과 같은 형식으로 작성하세요.
"""

# PREDEFINED_PROMPTS 수정
PREDEFINED_PROMPTS = {
    "1. 연구 목적": """
    사용자가 제공한 연구 주제와 키워드를 바탕으로, 연구 목적과 가설을 500자 이내의 줄글로 작성하세요. 어미는 반말 문어체로 합니다. (예: ~하였다. ~있다. ~있었다)
    다음 사항을 포함해야 합니다:
    1. 연구의 주요 목적
    2. 연구로 인해 의도하는 가설
    3. 가설을 입증하기 위한 구체적인 설명
    4. 연구의 중요성과 예상되는 결과
    5. 이 연구가 향후 진단과 치료에 있어 환자들에게 어떤 방식으로 이득이 될 것인지에 대한 설명

    사용자 입력:
    {user_input}

    위의 내용을 바탕으로 연구 목적과 가설을 구체화하여 작성해주세요. 특히, 연구 결과가 환자들에게 제공할 수 있는 실질적인 이익을 명확히 기술하여 IRB(기관 윤리위원회) 승인에 도움이 되도록 해주세요.
    """,
    
    "2. 연구 배경": """
    제공된 정보를 바탕으로 연구의 배경을 1000자 이내로 설명해주세요. 
    중요: 모든 문장은 반드시 반말 문어체로 끝나야 합니다. 예시: ~했다, ~였다, ~알려져 있다, ~보고되었다.
    
    다음 구조에 따라 연구 배경을 작성하세요. 각 부분을 별도의 문단으로 구성하고, 제시된 순서를 엄격히 지켜주세요. 단, 문단에 번호를 매기지 말고 자연스러운 줄글 형식으로 작성하세요:
    
    첫 번째 문단: 이론적 배경 및 근거
    - 연구 주제와 관련된 주요 이론과 개념을 설명한다.
    - 이 연구가 기반하는 이론적 틀을 제시한다.
    
    두 번째 문단: 선행 연구 및 결과, 국내외 연구 현황
    - 제공된 PDF 파일의 내용을 바탕으로 관련된 주요 선행 연구들을 요약한다.
    - 선행 연구의 주요 발견과 한계점을 언급한다.
    - 반드시 제공된 PDF 파일의 내용을 참고하여 설명한다. 
    - 인용 시 "한 연구에서는", "이전 연구에서는" 등의 표현을 사용하고, 각주 번호를 추가한다. 예: "한 연구에서는[1] ..."
    - 국내외 관련 연구의 현재 상태를 설명한다.
    - "국내"는 한국(Korea)을 의미한다.
    - "국외"는 한국(Korea)이 외 모든 나라를 의미한다.
    - 제공된 PDF 파일 내용 중 한국 소속 저자의 연구가 있다면, 이를 바탕으로 국내 연구 현황을 설명한다.
    
    세 번째 문단: 연구 배경과 연구의 정당성에 대한 설명
    - 현재 연구가 필요한 이유를 설명한다.
    - 이 연구가 어떤 점에서 중요한지 강조한다.
    - 본 연구가 기존 연구와 어떻게 다른지 또는 어떻게 기여할 수 있는지 설명한다.
    
    
    사용자 입력:
    {user_input}
    
    연구 목적:
    {research_purpose}
    
    PDF 내용:
    {pdf_content}
    (각 PDF 파일의 초록(abstract), 서론(introduction), 결론(conclusion) 섹션이 포함되어 있습니다.)

    한국 소속 저자 포함 여부:
    {korean_authors}
    
    위의 내용을 바탕으로 연구 배경을 구체화하여 작성해주세요. 특히 제공된 PDF 파일의 내용을 적극적으로 활용하여 연구 배경 작성에 참고해주세요. 참고문헌을 인용할 때는 [저자, 연도] 형식으로 표기해주세요.
    
    주의사항:
    1. 모든 문장은 반드시 '~다'로 끝나야 합니다. 존댓말 사용은 절대 금지입니다.
    2. 각 부분을 별도의 문단으로 작성하고, 제시된 순서를 반드시 지켜주세요.
    3. 문단에 번호를 매기지 말고, 자연스러운 줄글 형식으로 작성하세요.
    4. 각 문단의 내용이 명확히 구분되도록 작성해주세요.
    5. 두 번째 문단에서는 반드시 제공된 PDF 파일의 내용을 참고하여 설명해주세요.
    6. 두 번째 문단에서는 국내 연구 현황을 정확히 파악하여 설명해주세요. 한국 소속 저자의 연구가 있다면 반드시 포함시키세요.
    7. 선행 연구 내용은 제공된 PDF 내용(초록, 서론, 결론)만을 사용하여 작성하세요. 추가적인 정보를 임의로 만들어내지 마세요.
    8. 각 PDF 파일의 초록, 서론, 결론 내용을 적극적으로 활용하여 선행 연구를 요약하고 설명하세요.
    9. 사용자가 입력한 추가 정보나 키워드를 적절히 반영하여 연구 배경을 보완하세요.
    10. 두 번째 문단에서 절대로 제공된 PDF 내용에 없는 저자나 연구를 언급하지 마세요. 확실하지 않은 정보는 포함하지 마세요.
    11. 인용 시 저자 이름을 사용하지 말고, "한 연구에서는", "이전 연구에서는" 등의 표현을 사용하세요. 각 인용에 각주 번호를 추가하세요.
    12. 각주 번호는 참고문헌 목록의 순서와 일치해야 합니다.
    13. 제공된 정보와 선행 연구 내용을 적극적으로 활용하되, 연구 계획서 작성 과정에서 외부 자료를 참조했다는 언급은 피하세요. 예를 들어, "제공된 PDF에 따르면" 또는 "사용자가 제공한 정보에 의하면" 등의 표현은 사용하지 마세요. 대신 "선행 연구에 따르면" 또는 "기존 문헌에서는" 등의 표현을 사용하세요.
    
    위 지침을 엄격히 따라 연구 배경을 작성해주세요. 두 번째 문단에서 제공된 PDF 내용에 없는 정보는 절대 포함하지 마세요.
    """,

    "3. 선정기준, 제외기준": """
    1, 2번 섹션의 결과물과 참고한 논문들을 토대로, 이 연구에 적당한 대상자 그룹(선정기준)과 연구에서 제외해야 할 그룹(제외기준)을 추천해주세요. 다음 지침을 따라주세요:
    1. 구체적인 년도나 시기는 적지 않습니다. (잘못된 예시: 2009년 국가 건강검진을 받은 4,234,415명)
    2. 선정기준 예시: 40세에서 60세 사이에 해당하며, 이전 치매에 진단받은 과거력이 없는 수검자
    3. 제외기준 예시: 40세 이하 혹은 60세 이상, 검진 당시 치매 진단 과거력 있는 수검자, 누락된 변수 정보가 있는 수검자
    4. 이외 다른 말은 하지 말것.

    사용자 입력:
    {user_input}

    연구 목적:
    {research_purpose}

    연구 배경:
    {research_background}

    위의 내용을 바탕으로 적절한 선정기준, 제외기준을 작성해주세요.
    """,
    "4. 대상자 수 및 산출근거": """
    이전 섹션의 내용과 업로드된 논문들을 참고하여 다음 형식에 맞춰 대상자 수 및 산출근거를 작성해주세요. 어미는 반말 문어체로 합니다. (예: ~하였다. ~있다. ~있었다):

    다음 형식으로 대상자 수 및 산출근거를 작성해주세요:
    
    1) 대상자 수 (전체기관 대상자수/원내 대상자수)
    - 예상 연구대상자 수는 절대적이 아니며, 계획된 연구에서 필요한 결과를 얻을 수 있는 최소한 이상의 연구대상자 수
    [여기에 전체기관 대상자수]명/[여기에 원내 대상자수]명
    
    2) 산출 근거
    - 선행연구, 통계학적 평가방법에 근거하여 제시
    [여기에 산출 근거를 자세히 설명해주세요. 다음 사항을 포함하세요:]
    - 선행연구와 통계학적 평가방법에 근거한 설명
    - 가능한 경우, 구체적인 통계적 방법(예: 검정력 분석)을 언급하고 사용된 가정들을 설명
    - 대상자 수가 연구 목적을 달성하기에 충분한 이유를 설명
    
    주의사항:
    1. "사용자가 제시한", "제시된 대상자 수", "제안된" 등의 표현을 사용하지 마세요.
    2. 대상자 수를 연구 계획의 일부로 설명하고, 그 수를 선택한 이유를 객관적으로 서술하세요.
    3. 독자가 읽었을 때 대상자 수가 외부에서 입력되었다는 인상을 주지 않도록 주의하세요.
    4. 대상자 수가 연구 목적, 통계적 요구사항, 선행 연구 등을 고려하여 결정되었다는 점을 강조하세요.
    
    연구 목적:
    {research_purpose}
    
    연구 배경:
    {research_background}
    
    선정기준, 제외기준:
    {selection_criteria}
    
    사용자 입력 대상자 수 (참고용, 직접적으로 언급하지 마세요):
    - 총 대상자 수: {total_subjects}
    - 원내 대상자 수: {internal_subjects}
    - 타 기관 대상자 수: {external_subjects}
    
    위의 내용을 바탕으로 대상자 수 및 산출근거를 작성해주세요. 사용자 입력 대상자 수가 있다면 이를 참고하되, 직접적으로 언급하지 말고 연구 계획의 일부로 자연스럽게 설명해주세요. 입력된 대상자 수가 없다면, 연구 목적과 배경을 고려하여 적절한 대상자 수를 제안하고 그 근거를 설명해주세요.
    """,
    
    "5. 자료분석과 통계적 방법": """
    이전 섹션의 내용을 바탕으로 자료분석과 통계적 방법을 1000자 이내로 작성해주세요. 어미는 '-할 것이다', '-할 예정이다'와 같은 미래형 문어체로 작성합니다. 다음 사항을 포함해야 합니다:

    1. 수집해야 하는 수치나 값, 변수들 제시
    2. 변수의 이름은 영어로 작성하고, 긴 변수명의 경우 연구에 사용할 수 있는 약자도 함께 제시
    3. 연구에 사용할 군(group) 제시
    4. 연구를 통해 수집된 자료 또는 정보를 이용하는 방법(통계적 방법 포함) 기술
    5. 통계분석(계획) 제시:
       - 통계분석 방법
       - 분석대상군
       - 결측치의 처리 방법
       - 혼란변수 통제방법
       - 유의수준
       - 결과제시와 결과 도출 방안

    사용자 입력:
    {user_input}


    연구 목적:
    {research_purpose}

    연구 배경:
    {research_background}

    선정기준, 제외기준:
    {selection_criteria}

    대상자 수 및 산출근거:
    {sample_size}

    위의 내용을 바탕으로 자료분석과 통계적 방법을 구체적으로 작성해주세요. 각 항목을 명확히 구분하여 작성하되, 전체적으로 일관성 있는 내용이 되도록 해주세요.
    주의: 모든 설명은 미래형으로 작성해야 합니다. 예를 들어, "~~방법을 사용할 것이다", "~~를 분석할 예정이다" 등의 형식으로 작성하세요.
    """,
        
    "6. 연구방법": """
        1번부터 5번까지의 섹션 내용을 바탕으로 전체 연구방법을 500자 이내로 요약해주세요. 어미는 미래형 문어체로 통일합니다. (예: ~할 것이다. ~할 예정이다. ~할 계획이다) 다음 사항을 포함해야 합니다:
        
        1. 연구 목적의 핵심
        2. 연구 대상자 선정 및 제외 기준의 요점
        3. 대상자 수와 그 근거의 간략한 설명
        4. 주요 자료수집 방법
        5. 핵심적인 통계분석 방법
        
        이 연구가 어떤 방법으로 진행될 것인지 간단명료하게 설명해주세요. 전문적이면서도 이해하기 쉽게 작성해주세요. 모든 내용은 연구 계획에 대한 것이므로 미래형으로 작성해야 합니다.
        
        연구 목적:
        {research_purpose}
        
        연구 배경:
        {research_background}
        
        선정기준, 제외기준:
        {selection_criteria}
        
        대상자 수 및 산출근거:
        {sample_size}
        
        자료분석과 통계적 방법:
        {data_analysis}
        
        위의 내용을 바탕으로 전체 연구방법을 미래형으로 요약해주세요. 모든 문장이 미래형으로 작성되었는지 다시 한 번 확인하세요.
        """,
        
    "7. 연구 과제명": """
지금까지 작성된 연구계획서의 모든 내용을 바탕으로 연구 과제명을 추천해주세요. 다음 지침을 따라주세요:

1. 총 3가지의 제목 옵션을 제시해주세요.
2. 각 옵션은 영문 제목 1개와과 한글 제목 1개로 구성됩니다.
3. 영문 제목과 한글 제목은 별도의 줄에 작성합니다.
4. 각 옵션은 번호를 붙이지 않고 줄바꿈으로 구분합니다.
5. 제목은 완전한 문장으로, 잘리지 않게 작성합니다.
6. 제목은 연구의 핵심 내용을 간결하고 명확하게 표현해야 합니다.
7. 제목은 연구의 목적, 대상, 방법 등을 포함할 수 있습니다.
8. 영문 제목은 첫 글자만 대문자로 작성하세요. (예: Effect of...)
9. 제목 외 다른 말은 하지 마세요.

형식 예시:
[완전한 영문 제목]
[완전한 한글 제목]

[완전한 영문 제목]
[완전한 한글 제목]

[완전한 영문 제목]
[완전한 한글 제목]

사용자 입력:
{user_input}

연구 목적: {research_purpose}
연구 배경: {research_background}
선정기준, 제외기준: {selection_criteria}
대상자 수 및 산출근거: {sample_size}
자료분석과 통계적 방법: {data_analysis}
연구방법: {research_method}

위의 내용을 바탕으로 3가지의 연구 과제명 옵션을 제시해주세요.
"""
}


# 연구 섹션 순서 정의
RESEARCH_SECTIONS = [
    "1. 연구 목적",
    "2. 연구 배경",
    "3. 선정기준, 제외기준",
    "4. 대상자 수 및 산출근거",
    "5. 자료분석과 통계적 방법",
    "6. 연구방법",
    "7. 연구 과제명", 
    # 다른 섹션들은 나중에 추가할 예정입니다.
]

# Anthropic API 클라이언트 초기화 함수
def initialize_anthropic_client(api_key):
    try:
        client = anthropic.Client(api_key=api_key)
        # 간단한 API 호출로 키 유효성 검사
        client.messages.create(
            model="claude-3-5-sonnet-20241022",
            max_tokens=2000,
            messages=[{"role": "user", "content": "Hello"}]
        )
        return client
    except Exception as e:
        st.error(f"API 키 초기화 중 오류 발생: {str(e)}")
        return None

#세션 초기화 함수
def reset_session():
    for key in list(st.session_state.keys()):
        del st.session_state[key]
    st.session_state.clear()

                
#AI 응답 생성 함수
def generate_ai_response(prompt):
    if 'anthropic_client' in st.session_state and st.session_state.anthropic_client:
        try:
            system_prompt = f"{SYSTEM_PROMPT}\n\n추가 지시사항: 답변을 작성할 때 번호나 불렛 포인트를 사용하지 말고, 서술형으로 작성해주세요. 문단을 나누어 가독성 있게 작성하되, 전체적으로 하나의 연결된 글이 되도록 해주세요."
            
            response = st.session_state.anthropic_client.messages.create(
                model="claude-3-5-sonnet-20241022",
                max_tokens=2000,
                system=system_prompt,
                messages=[
                    {"role": "user", "content": prompt}
                ]
            )
            return response.content[0].text
        except anthropic.APIError as e:
            st.error(f"Anthropic API 오류: {str(e)}")
            return f"AI 응답 생성 중 API 오류가 발생했습니다: {str(e)}"
        except Exception as e:
            st.error(f"예상치 못한 오류 발생: {str(e)}")
            return f"AI 응답을 생성하는 중 예상치 못한 오류가 발생했습니다: {str(e)}"
    else:
        return "API 클라이언트가 초기화되지 않았습니다. API 키를 다시 확인해주세요."


# PDF 파일 업로드 함수
def upload_pdf():
    uploaded_file = st.file_uploader("PDF 파일을 업로드하세요.", type="pdf")
    if uploaded_file is not None:
        return extract_text_from_pdf(uploaded_file)
    return None

# PDF에서 텍스트 추출 함수
def extract_text_from_pdf(pdf_file):
    try:
        # pdfminer를 사용하여 텍스트 추출
        text = extract_text(pdf_file)
        return text
    except Exception as e:
        print(f"Error extracting text from {pdf_file.name}: {str(e)}")
        return ""

# Google Scholar 검색 함수 수정
def search_google_scholar(query, max_results=15):
    search_query = scholarly.search_pubs(query)
    results = defaultdict(list)
    keywords = query.lower().split()

    for result in search_query:
        if len(results['all_keywords']) + len(results['partial_keywords']) >= max_results:
            break
        try:
            # 논문 여부 확인 로직
            if not is_likely_paper(result):
                continue

            title = result['bib'].get('title', 'No title')
            year = result['bib'].get('pub_year', 'No year')
            authors = result['bib'].get('author', 'No author')
            if isinstance(authors, list):
                authors = ", ".join(authors[:2]) + "..." if len(authors) > 2 else ", ".join(authors)
            link = result.get('pub_url', '#')
            
            # 키워드 포함 여부 확인
            title_lower = title.lower()
            if all(keyword in title_lower for keyword in keywords):
                results['all_keywords'].append({"title": title, "year": year, "authors": authors, "link": link})
            elif any(keyword in title_lower for keyword in keywords):
                results['partial_keywords'].append({"title": title, "year": year, "authors": authors, "link": link})
        except AttributeError:
            continue

    # 각 그룹 내에서 최신 순으로 정렬
    for key in results:
        results[key].sort(key=lambda x: x['year'], reverse=True)

    # 모든 키워드 포함 결과를 먼저, 그 다음 부분 키워드 포함 결과 반환
    final_results = results['all_keywords'] + results['partial_keywords']
    return final_results[:max_results]

def is_likely_paper(result):
    # 논문일 가능성이 높은지 확인하는 함수
    bib = result.get('bib', {})
    
    # 1. 출판 유형 확인
    if bib.get('pub_type', '').lower() == 'book':
        return False
    
    # 2. 저널 정보 확인
    if 'journal' in bib or 'conference' in bib:
        return True
    
    # 3. 페이지 정보 확인 (논문은 대개 페이지 정보가 있음)
    if 'pages' in bib:
        return True
    
    # 4. 제목에 논문을 나타내는 키워드 확인
    title = bib.get('title', '').lower()
    paper_keywords = ['study', 'analysis', 'investigation', 'research', 'paper', 'article']
    if any(keyword in title for keyword in paper_keywords):
        return True
    
    # 5. 출판사 정보가 없으면 논문일 가능성이 높음 (책은 대개 출판사 정보가 있음)
    if 'publisher' not in bib:
        return True
    
    return False

# 참고문헌 정리 함수 추가
def format_references(scholar_results, pdf_files):
    references = []
    
    # Google Scholar 결과 처리
    for i, result in enumerate(scholar_results, start=len(references)+1):
        authors = result['authors'].split(', ')
        if len(authors) > 6:
            authors = authors[:6] + ['et al.']
        author_string = ', '.join(authors)
        reference = f"{i}. {author_string}. {result['title']} URL: {result['link']}."
        references.append(reference)
    
    # PDF 파일 처리
    for i, pdf_file in enumerate(pdf_files, start=len(references)+1):
        reference = f"{i}. {pdf_file.name}"
        references.append(reference)
    
    return references

# PDF에서 특정 섹션 추출하는 함수 
def extract_sections(text):
    sections = {
        'abstract': '',
        'introduction': '',
        'conclusion': ''
    }
    
    # Abstract 추출
    abstract_match = re.search(r'(?i)abstract.*?(?=\n\n|\n[A-Z])', text, re.DOTALL)
    if abstract_match:
        sections['abstract'] = abstract_match.group(0)
    
    # Introduction 추출
    intro_match = re.search(r'(?i)introduction.*?(?=\n\n|\n[A-Z])', text, re.DOTALL)
    if intro_match:
        sections['introduction'] = intro_match.group(0)
    
    # Conclusion 추출
    conclusion_match = re.search(r'(?i)conclusion.*?(?=\n\n|\n[A-Z]|$)', text, re.DOTALL)
    if conclusion_match:
        sections['conclusion'] = conclusion_match.group(0)
    
    return sections

# 1. 연구목적 작성 함수
def write_research_purpose():
    st.markdown("## 1. 연구 목적")
    
    # 히스토리 초기화
    if "1. 연구 목적_history" not in st.session_state:
        st.session_state["1. 연구 목적_history"] = []

    st.markdown("어떤 연구를 계획중인지, 연구에 대한 내용이나 키워드를 형식에 상관없이 자유롭게 입력해주세요.\n 기존에 작성한 초록이나 논문의 영어 원문이 있다면 붙여 넣어도 됩니다. \n입력 후 버튼을 누르면 AI 모델이 연구목적에 대한 줄글을 작성 해 줍니다.")
    
    user_input = st.text_area("연구 주제 및 키워드:", height=150)
    
    if st.button("연구 목적 AI 생성"):
        if user_input:
            prompt = PREDEFINED_PROMPTS["1. 연구 목적"].format(user_input=user_input)
            ai_response = generate_ai_response(prompt)
            
            # 현재 내용을 히스토리에 추가
            current_content = load_section_content("1. 연구 목적")
            if current_content:
                st.session_state["1. 연구 목적_history"].append(current_content)
            
            save_section_content("1. 연구 목적", ai_response)
            st.session_state.show_modification_request = False
            st.rerun()
        else:
            st.warning("연구 주제나 키워드를 입력해주세요.")

    # AI 응답 표시
    content = load_section_content("1. 연구 목적")
    if content:
        st.markdown("### AI가 생성한 연구 목적:")
        st.markdown(content)
        
        char_count = len(content)
        st.info(f"생성된 내용의 글자 수: {char_count}/500")
        
        if char_count > 500:
            st.warning("생성된 내용이 500자를 초과했습니다. 수정이 필요할 수 있습니다.")

        # 수정 요청 기능
        if st.button("수정 요청하기"):
            st.session_state.show_modification_request = True
            st.rerun()

        if st.session_state.get('show_modification_request', False):
            modification_request = st.text_area(
                "수정을 원하는 부분과 수정 방향을 설명해주세요:",
                height=150,
                key="modification_request_1"
            )
            if st.button("수정 요청 제출", key="submit_modification_1"):
                if modification_request:
                    current_content = load_section_content("1. 연구 목적")
                    # 현재 내용을 히스토리에 추가
                    st.session_state["1. 연구 목적_history"].append(current_content)
                    
                    prompt = f"""
                    현재 연구 목적:
                    {current_content}

                    사용자의 수정 요청:
                    {modification_request}

                    위의 수정 요청을 반영하여 연구 목적을 수정해주세요. 다음 지침을 따라주세요:
                    1. 전체 맥락을 유지하면서 내용을 수정하세요. 기존 내용을 완전히 그대로 유지하는 것이 아닙니다.
                    2. 수정 요청된 부분을 중점적으로 변경하되, 필요하다면 다른 부분도 조정하여 전체적인 일관성을 유지하세요.
                    3. 수정된 내용은 자연스럽게 기존 내용과 연결되어야 합니다.
                    4. 전체 내용은 500자를 넘지 않아야 합니다.
                    5. 수정된 부분은 기존 내용의 맥락과 일관성을 유지해야 합니다.
                    7. 어미는 반말 문어체로 합니다. (예: ~하였다. ~있다. ~있었다)
                    8. 각 부분의 길이를 조절하여 전체 글자 수 제한을 지키세요.
                    8. 내용 이외 다른말은 하지 마세요.
                    
                    수정된 전체 연구 목적을 작성해주세요.
                    """
                    modified_response = generate_ai_response(prompt)
                    
                    save_section_content("1. 연구 목적", modified_response)
                    st.session_state.show_modification_request = False
                    st.rerun()
                else:
                    st.warning("수정 요청 내용을 입력해주세요.")

    # 편집 기능
    edited_content = st.text_area(
        "생성된 내용을 편집하거나 내용을 직접 입력하세요:",
        content,
        height=200,
        key="edit_content_1"
    )
    st.warning("다음 섹션으로 넘어가기 전에 편집내용 저장 버튼을 누르세요.")
    col1, col2 = st.columns(2)
    with col1:
        if st.button("편집 내용 저장", key="save_edit_1"):
            # 현재 내용을 히스토리에 추가
            st.session_state["1. 연구 목적_history"].append(content)
            save_section_content("1. 연구 목적", edited_content)
            st.success("편집된 내용이 저장되었습니다.")
            st.rerun()
    with col2:
        if st.button("실행 취소", key="undo_edit_1"):
            if st.session_state["1. 연구 목적_history"]:
                # 히스토리에서 마지막 항목을 가져와 현재 내용으로 설정
                previous_content = st.session_state["1. 연구 목적_history"].pop()
                save_section_content("1. 연구 목적", previous_content)
                st.success("이전 버전으로 되돌렸습니다.")
                st.rerun()
            else:
                st.warning("더 이상 되돌릴 수 있는 버전이 없습니다.")


# 2. 연구 배경 작성 함수
def write_research_background():
    st.markdown("## 2. 연구 배경")

    # 히스토리 초기화
    if "2. 연구 배경_history" not in st.session_state:
        st.session_state["2. 연구 배경_history"] = []

    # 사용자 입력 받기
    user_input = st.text_area("연구 배경에 대해 작성하고 싶은 내용이나 AI 모델이 참고 해야할 내용이 있다면 입력해주세요. 없으면 빈칸으로 두고 진행해도 됩니다. \n빈칸으로 둘 경우 `2.연구목적` 섹션의 내용과 업로드된 PDF를 기준으로 내용을 작성합니다.:", height=150)

     # 참조논문 검색 부분을 expander로 감싸기
    with st.expander("참조논문 검색하기", expanded=False):
        # 키워드 입력
        keywords = st.text_input("연구 배경 작성을 위한 참조논문 검색에 사용할 키워드를 입력하세요 (최대 10개, 쉼표로 구분):")
        keywords_list = [k.strip() for k in keywords.split(',') if k.strip()][:10]
        
        if keywords_list:
            st.write("입력된 키워드:", ", ".join(keywords_list))
            
        if st.button("논문 검색"):
            if keywords_list:
                search_query = " ".join(keywords_list)
                
                with st.spinner("논문을 검색 중입니다..."):
                    scholar_results = search_google_scholar(search_query)
                    
                st.session_state.scholar_results = scholar_results
                st.success("검색이 완료되었습니다.")
                st.rerun()
                
        # 검색 결과 표시
        if 'scholar_results' in st.session_state:
            st.subheader("Google Scholar 검색 결과 (최대 15개)")
            for i, result in enumerate(st.session_state.scholar_results):
                col1, col2 = st.columns([3, 1])
                with col1:
                    st.markdown(f"[{result['title']} ({result['year']})]({result['link']})")
                    st.caption(f"저자: {result['authors']}")
                with col2:
                    if st.button("삭제", key=f"del_scholar_{i}"):
                        del st.session_state.scholar_results[i]
                        st.rerun()

    # 새로운 텍스트 추가
    st.markdown("""
    검색한 논문을 내용을 쉽게 한글 요약해서 보시려면 "병리 논문 요약하기📝 ver.2 (HJY)" 을 사용해보세요! [링크](https://journalsummaryver2.streamlit.app/)
    """)
    
    # PDF 파일 업로드 
    uploaded_files = st.file_uploader("연구 배경 작성에 참고할 선행연구 논문 PDF 파일을 업로드하세요. 중요한 논문 위주로 4개 이하 업로드를 추천합니다. \n**주의:** 검색 결과의 논문 내용은 자동으로 반영되지 않습니다. \n검색된 논문들을 사용하시려면 각 웹페이지에서 PDF 파일을 다운 받은 후 여기에 업로드 하세요.", type="pdf", accept_multiple_files=True)
    
    if uploaded_files:
        st.session_state.pdf_texts = []
        st.session_state.pdf_files = uploaded_files
        st.session_state.pdf_metadata = []
        for uploaded_file in uploaded_files:
            pdf_text = extract_text_from_pdf(uploaded_file)
            st.session_state.pdf_texts.append(pdf_text)
            metadata = extract_references(pdf_text)
            st.session_state.pdf_metadata.append(metadata)
        st.success(f"{len(uploaded_files)}개의 PDF 파일이 성공적으로 업로드되었습니다.")


    # 연구 배경 생성 버튼
    if st.button("연구배경 AI 생성 요청하기"):
        if 'pdf_texts' in st.session_state and st.session_state['pdf_texts']:
            research_purpose = load_section_content("1. 연구 목적")
            
            pdf_contents = []
            korean_authors = False
            for i, pdf_text in enumerate(st.session_state['pdf_texts']):
                extracted_sections = extract_sections(pdf_text)
                metadata = st.session_state.get('pdf_metadata', [])
                if i < len(metadata):
                    current_metadata = metadata[i]
                    if isinstance(current_metadata, dict):
                        is_korean = current_metadata.get('is_korean', False)
                    else:
                        is_korean = False
                else:
                    is_korean = False

                pdf_contents.append({
                    "file_name": st.session_state['pdf_files'][i].name,
                    "abstract": extracted_sections['abstract'],
                    "introduction": extracted_sections['introduction'],
                    "conclusion": extracted_sections['conclusion'],
                    "is_korean": is_korean
                })
                if is_korean:
                    korean_authors = True
            
            pdf_content_json = json.dumps(pdf_contents)
            
            prompt = PREDEFINED_PROMPTS["2. 연구 배경"].format(
                user_input=user_input,
                keywords=keywords,
                research_purpose=research_purpose,
                pdf_content=pdf_content_json,
                korean_authors=korean_authors
            )
            
            # 추출된 참고문헌 정보 추가
            prompt += "\n\n다음은 제공된 PDF 파일들의 정확한 참고문헌 정보입니다. 연구 배경 작성 시 반드시 이 정보만을 사용하여 인용해주세요:\n"
            for metadata in st.session_state.pdf_metadata:
                if metadata:  # metadata가 비어있지 않은 경우에만 처리
                    author = metadata[0][0] if metadata[0] else "Unknown"
                    year = metadata[0][1] if len(metadata[0]) > 1 else "Unknown"
                    prompt += f"[{author}, {year}]\n"
            
            ai_response = generate_ai_response(prompt)

            # AI 응답 검증 및 수정
            verified_response = verify_and_correct_references(ai_response, st.session_state.pdf_metadata)
            
            save_section_content("2. 연구 배경", verified_response)
            
            # 현재 내용을 히스토리에 추가
            current_content = load_section_content("2. 연구 배경")
            if current_content:
                st.session_state["2. 연구 배경_history"].append(current_content)
            
            save_section_content("2. 연구 배경", ai_response)
            st.session_state.show_modification_request_2 = False
            st.rerun()
        else:
            st.warning("PDF를 업로드한 후 다시 시도해주세요.")

    # AI 응답 표시
    content = load_section_content("2. 연구 배경")
    if content:
        st.markdown("### AI가 생성한 연구 배경 (1000자 내외):")
        st.markdown(content)
        
        char_count = len(content)
        st.info(f"생성된 내용의 글자 수: {char_count}/1000")
        
        if char_count > 1000:
            st.warning("생성된 내용이 1000자를 초과했습니다. 수정이 필요할 수 있습니다.")

        # 수정 요청 기능
        if st.button("수정 요청하기", key="request_modification_2"):
            st.session_state.show_modification_request_2 = True
            st.rerun()

        if st.session_state.get('show_modification_request_2', False):
            modification_request = st.text_area(
                "수정을 원하는 부분과 수정 방향을 설명해주세요:",
                height=150,
                key="modification_request_2"
            )
            if st.button("수정 요청 제출", key="submit_modification_2"):
                if modification_request:
                    current_content = load_section_content("2. 연구 배경")
                    # 현재 내용을 히스토리에 추가
                    st.session_state["2. 연구 배경_history"].append(current_content)
                    
                    prompt = f"""
                    현재 연구 배경:
                    {current_content}

                    사용자의 수정 요청:
                    {modification_request}

                   위의 수정 요청을 반영하여 연구 배경을 수정해주세요. 다음 지침을 엄격히 따라주세요:

                    1. 전체 맥락을 유지하면서 내용을 수정하세요. 기존 내용을 그대로 유지하는 것이 아닙니다.
                    2. 수정 요청된 부분을 중점적으로 변경하되, 필요하다면 다른 부분도 조정하여 전체적인 일관성을 유지하세요.
                    3. 수정된 내용은 자연스럽게 기존 맥락과 연결되어야 합니다.
                    4. 전체 내용은 반드시 1000자 이내로 작성하세요. 이는 엄격한 제한사항입니다.
                    5. 연구 배경의 논리적 흐름을 유지하세요.
                    6. 어미는 반말 문어체로 통일하세요. (예: ~하였다. ~있다. ~있었다)
                    7. 다음 구조를 유지하되, 각 부분의 길이를 조절하여 전체 글자 수 제한을 지키세요:
                       - 이론적 배경 및 근거
                       - 선행 연구 및 결과, 국내외 연구 현황 (검색된 논문 참고)
                       - 연구 배경과 연구의 정당성
                    8. 내용 이외의 다른 말은 하지 마세요.
                    
                    수정된 전체 연구 배경을 1000자 이내로 작성해주세요.
                    """
                    modified_response = generate_ai_response(prompt)
                    
                    save_section_content("2. 연구 배경", modified_response)
                    st.session_state.show_modification_request_2 = False
                    st.rerun()
                else:
                    st.warning("수정 요청 내용을 입력해주세요.")

    # 편집 기능
    edited_content = st.text_area(
        "생성된 내용을 편집하세요 :",
        content,
        height=300,
        key="edit_content_2"
    )
    st.warning("다음 섹션으로 넘어가기 전에 편집내용 저장 버튼을 누르세요.")
    col1, col2 = st.columns(2)
    with col1:
        if st.button("편집 내용 저장", key="save_edit_2"):
            # 현재 내용을 히스토리에 추가
            st.session_state["2. 연구 배경_history"].append(content)
            save_section_content("2. 연구 배경", edited_content)
            st.success("편집된 내용이 저장되었습니다.")
            st.rerun()
    with col2:
        if st.button("실행 취소", key="undo_edit_2"):
            if st.session_state["2. 연구 배경_history"]:
                # 히스토리에서 마지막 항목을 가져와 현재 내용으로 설정
                previous_content = st.session_state["2. 연구 배경_history"].pop()
                save_section_content("2. 연구 배경", previous_content)
                st.success("이전 버전으로 되돌렸습니다.")
                st.rerun()
            else:
                st.warning("더 이상 되돌릴 수 있는 버전이 없습니다.")

def verify_and_correct_references(response, correct_metadata):
    # 응답에서 참고문헌 추출
    cited_references = extract_references(response)
    
    # 추출된 참고문헌과 원본 메타데이터 비교 및 수정
    for ref in cited_references:
        ref_str = ', '.join(ref) if isinstance(ref, (list, tuple)) else str(ref)
        if ref_str not in [', '.join(map(str, m)) for m in correct_metadata]:
            # 잘못된 참고문헌 찾아 수정
            correct_ref = find_closest_match(ref_str, correct_metadata)
            response = response.replace(ref_str, ', '.join(map(str, correct_ref)))
    
    return response

def find_closest_match(ref, correct_metadata):
    def format_metadata(x):
        if isinstance(x, (list, tuple)):
            return ', '.join(str(item) for item in x)
        return str(x)
    
    return max(correct_metadata, key=lambda x: similarity(ref, format_metadata(x)))[0]

def similarity(a, b):
    # 간단한 유사도 계산 (예: 레벤슈타인 거리 사용)
    return 1 - (levenshtein_distance(a, b) / max(len(a), len(b)))

def levenshtein_distance(s1, s2):
    if len(s1) < len(s2):
        return levenshtein_distance(s2, s1)
    if len(s2) == 0:
        return len(s1)
    previous_row = range(len(s2) + 1)
    for i, c1 in enumerate(s1):
        current_row = [i + 1]
        for j, c2 in enumerate(s2):
            insertions = previous_row[j + 1] + 1
            deletions = current_row[j] + 1
            substitutions = previous_row[j] + (c1 != c2)
            current_row.append(min(insertions, deletions, substitutions))
        previous_row = current_row
    return previous_row[-1]

# 3. 선정기준, 제외기준 작성 함수
def write_selection_criteria():
    st.markdown("## 3. 선정기준, 제외기준")
    
    # 히스토리 초기화
    if "3. 선정기준, 제외기준_history" not in st.session_state:
        st.session_state["3. 선정기준, 제외기준_history"] = []

    # 사용자 입력 받기
    user_input = st.text_area("선정기준과 제외기준에 대해 AI에게 알려 줄 추가 정보나 고려사항이 있다면 입력해주세요.\n 특별히 없다면 빈칸으로 두어도 됩니다. 빈칸이라면 이전 섹션들의 내용을 기반으로 선정기준, 제외기준을 제안합니다:", height=150)

    if st.button("선정, 제외기준 AI에게 추천받기"):
        research_purpose = load_section_content("1. 연구 목적")
        research_background = load_section_content("2. 연구 배경")
        
        prompt = PREDEFINED_PROMPTS["3. 선정기준, 제외기준"].format(
            user_input=user_input,
            research_purpose=research_purpose,
            research_background=research_background
        )
        
        ai_response = generate_ai_response(prompt)
        
        # 현재 내용을 히스토리에 추가
        current_content = load_section_content("3. 선정기준, 제외기준")
        if current_content:
            st.session_state["3. 선정기준, 제외기준_history"].append(current_content)
        
        save_section_content("3. 선정기준, 제외기준", ai_response)
        st.rerun()

    # AI 응답 표시
    content = load_section_content("3. 선정기준, 제외기준")
    if content:
        st.markdown("### AI가 추천한 선정, 제외기준:")
        st.markdown(content)

        # 수정 요청 기능
        if st.button("수정 요청하기", key="request_modification_3"):
            st.session_state.show_modification_request_3 = True
            st.rerun()

        if st.session_state.get('show_modification_request_3', False):
            modification_request = st.text_area(
                "수정을 원하는 부분과 수정 방향을 설명해주세요:",
                height=150,
                key="modification_request_3"
            )
            if st.button("수정 요청 제출", key="submit_modification_3"):
                if modification_request:
                    current_content = load_section_content("3. 선정기준, 제외기준")
                    # 현재 내용을 히스토리에 추가
                    st.session_state["3. 선정기준, 제외기준_history"].append(current_content)
                    
                    prompt = f"""
                    현재 선정기준, 제외기준:
                    {current_content}

                    사용자의 수정 요청:
                    {modification_request}

                    위의 수정 요청을 반영하여 선정기준, 제외기준을 수정해주세요. 어미는 문어제 반말을 사용하세요.(예시: "~했다.", "~있다.", "~이다.") 다음 지침을 따라주세요:
                    1. 전체 맥락을 유지하면서 내용을 수정하세요. 기존 내용을 그대로 유지하는 것이 아닙니다.
                    2. 수정 요청된 부분을 중점적으로 변경하되, 필요하다면 다른 부분도 조정하여 전체적인 일관성을 유지하세요.
                    3. 수정된 내용은 자연스럽게 기존 맥락과 연결되어야 합니다.
                    4. 수정된 내용은 자연스럽게 기존 내용과 연결되어야 합니다.
                    5. 수정된 부분은 기존 내용의 맥락과 일관성을 유지해야 합니다.
                    6. 내용 이외 다른말은 하지 말것.
                    
                    수정된 전체 선정기준, 제외기준을 작성해주세요.
                    """
                    modified_response = generate_ai_response(prompt)
                    
                    save_section_content("3. 선정기준, 제외기준", modified_response)
                    st.session_state.show_modification_request_3 = False
                    st.rerun()
                else:
                    st.warning("수정 요청 내용을 입력해주세요.")
    
    # 편집 기능
    edited_content = st.text_area(
        "생성된 내용을 편집하세요:",
        content,
        height=200,
        key="edit_content_3"
    )

    st.warning("다음 섹션으로 넘어가기 전에 편집내용 저장 버튼을 누르세요.")
    col1, col2 = st.columns(2)
    with col1:
        if st.button("편집 내용 저장", key="save_edit_3"):
            # 현재 내용을 히스토리에 추가
            st.session_state["3. 선정기준, 제외기준_history"].append(content)
            save_section_content("3. 선정기준, 제외기준", edited_content)
            st.success("편집된 내용이 저장되었습니다.")
            st.rerun()
    with col2:
        if st.button("실행 취소", key="undo_edit_3"):
            if st.session_state["4. 선정기준, 제외기준_history"]:
                # 히스토리에서 마지막 항목을 가져와 현재 내용으로 설정
                previous_content = st.session_state["3. 선정기준, 제외기준_history"].pop()
                save_section_content("3. 선정기준, 제외기준", previous_content)
                st.success("이전 버전으로 되돌렸습니다.")
                st.rerun()
            else:
                st.warning("더 이상 되돌릴 수 있는 버전이 없습니다.")

# 4. 대상자 수 및 산출근거 작성 함수 (수정)
def write_sample_size():
    st.markdown("## 4. 대상자 수 및 산출근거")
    
    # 히스토리 초기화
    if "4. 대상자 수 및 산출근거_history" not in st.session_state:
        st.session_state["4. 대상자 수 및 산출근거_history"] = []

    # 안내 텍스트 추가
    st.write("이미 정해진 대상자 수가 있다면 입력해주세요. \n없다면 비워두고 AI 추천받기 버튼을 눌러 추천받으세요.")

        # 사용자 입력 대상자 수
    col1, col2 = st.columns(2)
    with col1:
        internal_subjects = st.number_input("원내 대상자 수", min_value=0, value=None, step=1)
    with col2:
        external_subjects = st.number_input("타 기관 대상자 수", min_value=0, value=None, step=1)
    
    # 입력값이 None인 경우를 처리
    if internal_subjects is not None and external_subjects is not None:
        total_subjects = internal_subjects + external_subjects
        st.write(f"총 대상자 수: {total_subjects}명")
    else:
        st.write("대상자 수가 입력되지 않았습니다. AI에게 추천을 받으세요.")

    if st.button("대상자 수 및 산출근거 AI에게 추천받기"):
        research_purpose = load_section_content("1. 연구 목적")
        research_background = load_section_content("2. 연구 배경")
        selection_criteria = load_section_content("3. 선정기준, 제외기준")
        
        prompt = PREDEFINED_PROMPTS["4. 대상자 수 및 산출근거"].format(
            research_purpose=research_purpose,
            research_background=research_background,
            selection_criteria=selection_criteria,
            total_subjects=total_subjects if internal_subjects is not None and external_subjects is not None else "미입력",
            internal_subjects=internal_subjects if internal_subjects is not None else "미입력",
            external_subjects=external_subjects if external_subjects is not None else "미입력"
        )
        
        ai_response = generate_ai_response(prompt)
        
       # 현재 내용을 히스토리에 추가
        current_content = load_section_content("4. 대상자 수 및 산출근거")
        if current_content:
            st.session_state["4. 대상자 수 및 산출근거_history"].append(current_content)
        
        save_section_content("4. 대상자 수 및 산출근거", ai_response)
        st.rerun()

    # AI 응답 표시
    content = load_section_content("4. 대상자 수 및 산출근거")
    if content:
        st.markdown("### AI가 추천한 대상자 수 및 산출근거:")
        st.markdown(content)

        # 수정 요청 기능
        if st.button("수정 요청하기", key="request_modification_4"):
            st.session_state.show_modification_request_4 = True
            st.rerun()

        if st.session_state.get('show_modification_request_4', False):
            modification_request = st.text_area(
                "수정을 원하는 부분과 수정 방향을 설명해주세요:",
                height=150,
                key="modification_request_4"
            )
            if st.button("수정 요청 제출", key="submit_modification_4"):
                if modification_request:
                    current_content = load_section_content("4. 대상자 수 및 산출근거")
                    # 현재 내용을 히스토리에 추가
                    st.session_state["4. 대상자 수 및 산출근거_history"].append(current_content)
                    
                    prompt = f"""
                    현재 대상자 수 및 산출근거:
                    {current_content}

                    사용자의 수정 요청:
                    {modification_request}

                    위의 수정 요청을 반영하여 대상자 수 및 산출근거를 수정해주세요. 다음 지침을 따라주세요:
                    1. 전체 맥락을 유지하면서 내용을 수정하세요. 기존 내용을 그대로 유지하는 것이 아닙니다.
                    2. 수정 요청된 부분을 중점적으로 변경하되, 필요하다면 다른 부분도 조정하여 전체적인 일관성을 유지하세요.
                    3. 수정된 내용은 자연스럽게 기존 맥락과 연결되어야 합니다.
                    4. 수정된 부분은 기존 내용의 맥락과 일관성을 유지해야 합니다.
                    5. 어미는 반말 문어체로 합니다. (예: ~하였다. ~있다. ~있었다)
                    6. 내용 이외 다른말은 하지 말것.
                    
                    수정된 전체 대상자 수 및 산출근거를 작성해주세요.
                    """
                    modified_response = generate_ai_response(prompt)
                    
                    save_section_content("4. 대상자 수 및 산출근거", modified_response)
                    st.session_state.show_modification_request_4 = False
                    st.rerun()
                else:
                    st.warning("수정 요청 내용을 입력해주세요.")
    
    # 편집 기능
    edited_content = st.text_area(
        "생성된 내용을 편집하세요:",
        content,
        height=300,
        key="edit_content_4"
    )

    st.warning("다음 섹션으로 넘어가기 전에 편집내용 저장 버튼을 누르세요.")
    col1, col2 = st.columns(2)
    with col1:
        if st.button("편집 내용 저장", key="save_edit_4"):
            # 현재 내용을 히스토리에 추가
            st.session_state["4. 대상자 수 및 산출근거_history"].append(content)
            save_section_content("4. 대상자 수 및 산출근거", edited_content)
            st.success("편집된 내용이 저장되었습니다.")
            st.rerun()
    with col2:
        if st.button("실행 취소", key="undo_edit_4"):
            if st.session_state["4. 대상자 수 및 산출근거_history"]:
                # 히스토리에서 마지막 항목을 가져와 현재 내용으로 설정
                previous_content = st.session_state["4. 대상자 수 및 산출근거_history"].pop()
                save_section_content("4. 대상자 수 및 산출근거", previous_content)
                st.success("이전 버전으로 되돌렸습니다.")
                st.rerun()
            else:
                st.warning("더 이상 되돌릴 수 있는 버전이 없습니다.")

#5. 자료분석과 통계적 방법 함수
def write_data_analysis():
    st.markdown("## 5. 자료분석과 통계적 방법")
    
    # 히스토리 초기화
    if "5. 자료분석과 통계적 방법_history" not in st.session_state:
        st.session_state["5. 자료분석과 통계적 방법_history"] = []

    # 사용자 입력 받기
    user_input = st.text_area("자료분석과 통계적 방법에 대해 AI에게 알려줄 추가 정보나 고려사항이 있다면 입력해주세요. 특별히 없다면 빈칸으로 두어도 됩니다. \n빈칸이라면 이전 섹션들의 내용을 기반으로 선정기준, 제외기준을 제안합니다:", height=150)

    if st.button("자료분석 및 통계방법 AI에게 추천받기"):
        research_purpose = load_section_content("1. 연구 목적")
        research_background = load_section_content("2. 연구 배경")
        selection_criteria = load_section_content("3. 선정기준, 제외기준")
        sample_size = load_section_content("4. 대상자 수 및 산출근거")
        
        prompt = PREDEFINED_PROMPTS["5. 자료분석과 통계적 방법"].format(
            user_input=user_input,
            research_purpose=research_purpose,
            research_background=research_background,
            selection_criteria=selection_criteria,
            sample_size=sample_size
        )
        
        ai_response = generate_ai_response(prompt)
        
        # 현재 내용을 히스토리에 추가
        current_content = load_section_content("5. 자료분석과 통계적 방법")
        if current_content:
            st.session_state["5. 자료분석과 통계적 방법_history"].append(current_content)
        
        save_section_content("5. 자료분석과 통계적 방법", ai_response)
        st.rerun()

    # AI 응답 표시
    content = load_section_content("5. 자료분석과 통계적 방법")
    if content:
        st.markdown("### AI가 추천한 자료분석과 통계적 방법 (1000자 내외):")
        st.markdown(content)

        # 수정 요청 기능
        if st.button("수정 요청하기", key="request_modification_5"):
            st.session_state.show_modification_request_5 = True
            st.rerun()

        if st.session_state.get('show_modification_request_5', False):
            modification_request = st.text_area(
                "수정을 원하는 부분과 수정 방향을 설명해주세요:",
                height=150,
                key="modification_request_5"
            )
            if st.button("수정 요청 제출", key="submit_modification_5"):
                if modification_request:
                    current_content = load_section_content("5. 자료분석과 통계적 방법")
                    # 현재 내용을 히스토리에 추가
                    st.session_state["5. 자료분석과 통계적 방법_history"].append(current_content)
                    
                    prompt = f"""
                    현재 자료분석과 통계적 방법:
                    {current_content}

                    사용자의 수정 요청:
                    {modification_request}

                    위의 수정 요청을 반영하여 자료분석과 통계적 방법을 수정해주세요. 다음 지침을 따라주세요:
                    1. 모든 내용은 반드시 미래형으로 작성해야 합니다. 예시:
                       - "~~방법을 사용할 것이다", "~~를 분석할 예정이다", "~~을 수행할 계획이다" 등
                       - 과거형이나 현재형 표현(예: "분석하였다", "평가한다")은 절대 사용하지 마세요.
                    
                    2. 전체 맥락을 유지하면서 내용을 수정하세요. 기존 내용을 그대로 유지하는 것이 아닙니다.
                    3. 수정 요청된 부분을 중점적으로 변경하되, 필요하다면 다른 부분도 조정하여 전체적인 일관성을 유지하세요.
                    4. 수정된 내용은 자연스럽게 기존 맥락과 연결되어야 합니다.
                    5. 전체 내용은 1000자를 넘지 않아야 합니다.
                    6. 연구 방법의 논리적 흐름을 유지하세요.
                    7. 내용 이외의 다른 말은 하지 마세요.

                    수정된 전체 자료분석과 통계적 방법을 작성해주세요. 모든 문장이 미래형으로 작성되었는지 다시 한 번 확인하세요.
                    """
                    modified_response = generate_ai_response(prompt)
                    
                    save_section_content("6. 자료분석과 통계적 방법", modified_response)
                    st.session_state.show_modification_request_5 = False
                    st.rerun()
                else:
                    st.warning("수정 요청 내용을 입력해주세요.")
    
    # 편집 기능
    edited_content = st.text_area(
        "생성된 내용을 편집하세요:",
        content,
        height=300,
        key="edit_content_5"
    )

    st.warning("다음 섹션으로 넘어가기 전에 편집내용 저장 버튼을 누르세요.")
    col1, col2 = st.columns(2)
    with col1:
        if st.button("편집 내용 저장", key="save_edit_5"):
            # 현재 내용을 히스토리에 추가
            st.session_state["5. 자료분석과 통계적 방법_history"].append(content)
            save_section_content("5. 자료분석과 통계적 방법", edited_content)
            st.success("편집된 내용이 저장되었습니다.")
            st.rerun()
    with col2:
        if st.button("실행 취소", key="undo_edit_5"):
            if st.session_state["5. 자료분석과 통계적 방법_history"]:
                # 히스토리에서 마지막 항목을 가져와 현재 내용으로 설정
                previous_content = st.session_state["5. 자료분석과 통계적 방법_history"].pop()
                save_section_content("5. 자료분석과 통계적 방법", previous_content)
                st.success("이전 버전으로 되돌렸습니다.")
                st.rerun()
            else:
                st.warning("더 이상 되돌릴 수 있는 버전이 없습니다.")

    # 글자 수 표시
    if content:
        char_count = len(content)
        st.info(f"현재 글자 수: {char_count}/1000")
        if char_count > 1000:
            st.warning("글자 수가 1000자를 초과했습니다. 내용을 줄여주세요.")

#6. 연구방법 정리 함수
def write_research_method():
    st.markdown("## 6. 연구방법")
        
    # 안내 글 추가
    st.markdown("""
    `연구 방법`은 앞서 작성한 모든 섹션의 내용을 약 500자 가량의 글로 요약 정리해서 연구의 방법을 한눈에 볼 수 있도록 제시합니다. 

    AI의 내용 정리를 받기 원하면 아래 버튼을 누르고, 직접 작성하시려면 버튼아래 편집창에 직접 작성도 가능합니다.
    """)
    
    # 히스토리 초기화
    if "6. 연구방법_history" not in st.session_state:
        st.session_state["6. 연구방법_history"] = []

    if st.button("연구방법 정리 요청하기"):
        research_purpose = load_section_content("1. 연구 목적")
        research_background = load_section_content("2. 연구 배경")
        selection_criteria = load_section_content("3. 선정기준, 제외기준")
        sample_size = load_section_content("4. 대상자 수 및 산출근거")
        data_analysis = load_section_content("5. 자료분석과 통계적 방법")
        
        prompt = PREDEFINED_PROMPTS["6. 연구방법"].format(
            research_purpose=research_purpose,
            research_background=research_background,
            selection_criteria=selection_criteria,
            sample_size=sample_size,
            data_analysis=data_analysis
        )
        
        ai_response = generate_ai_response(prompt)
        
        # 현재 내용을 히스토리에 추가
        current_content = load_section_content("6. 연구방법")
        if current_content:
            st.session_state["6. 연구방법_history"].append(current_content)
        
        save_section_content("6. 연구방법", ai_response)
        st.rerun()

    # AI 응답 표시
    content = load_section_content("6. 연구방법")
    if content:
        st.markdown("### AI가 정리한 연구방법:")
        st.markdown(content)

        # 수정 요청 기능
        if st.button("수정 요청하기", key="request_modification_6"):
            st.session_state.show_modification_request_6 = True
            st.rerun()

        if st.session_state.get('show_modification_request_6', False):
            modification_request = st.text_area(
                "수정을 원하는 부분과 수정 방향을 설명해주세요:",
                height=150,
                key="modification_request_6"
            )
            if st.button("수정 요청 제출", key="submit_modification_6"):
                if modification_request:
                    current_content = load_section_content("6. 연구방법")
                    # 현재 내용을 히스토리에 추가
                    st.session_state["6. 연구방법_history"].append(current_content)
                    
                    prompt = f"""
                    현재 연구방법:
                    {current_content}

                    사용자의 수정 요청:
                    {modification_request}

                    위의 수정 요청을 반영하여 연구방법을 수정해주세요. 다음 지침을 따라주세요:
                    1. 모든 내용은 반드시 미래형으로 작성해야 합니다. 예시:
                       - "~~방법을 사용할 것이다", "~~를 분석할 예정이다", "~~을 수행할 계획이다" 등
                       - 과거형이나 현재형 표현(예: "분석하였다", "평가한다")은 절대 사용하지 마세요.
                    2. 사용자의 수정 요청을 최우선으로 고려하여 반영하세요. 수정 요청의 각 항목을 하나씩 확인하고, 모두 반영되었는지 확인하세요.
                    3. 전체 맥락을 유지하면서 내용을 수정하세요. 기존 내용을 그대로 유지하는 것이 아닙니다.
                    4. 수정된 내용은 자연스럽게 기존 맥락과 연결되어야 합니다.
                    5. 전체 내용은 500자를 넘지 않아야 합니다.
                    6. 연구 방법의 논리적 흐름을 유지하세요.
                    7. 어미는 "-할 것이다", "-할 예정이다", "-할 계획이다" 등의 미래형 문어체로 통일하세요.
                    8. 내용 이외의 다른 말은 하지 마세요.
                    9. 수정 작업 후, 다음을 확인하세요:
                       a) 사용자의 모든 수정 요청이 반영되었는지
                       b) 모든 문장이 미래형으로 작성되었는지
                       c) 전체적인 내용의 일관성과 논리적 흐름이 유지되는지
                    
                    수정된 전체 연구방법을 작성해주세요. 모든 문장이 미래형으로 작성되었는지 다시 한 번 확인하세요.
                    """
                    modified_response = generate_ai_response(prompt)
                    
                    save_section_content("6. 연구방법", modified_response)
                    st.session_state.show_modification_request_6 = False
                    st.rerun()
                else:
                    st.warning("수정 요청 내용을 입력해주세요.")
    
    # 편집 기능
    edited_content = st.text_area(
        "연구방법을 직접 여기에 작성하거나, 위 버튼을 눌러 AI의 정리를 받으세요. 생성된 내용을 편집하세요:",
        content,
        height=150,
        key="edit_content_6"
    )

    st.warning("다음 섹션으로 넘어가기 전에 편집내용 저장 버튼을 누르세요.")
    col1, col2 = st.columns(2)
    with col1:
        if st.button("편집 내용 저장", key="save_edit_6"):
            # 현재 내용을 히스토리에 추가
            st.session_state["6. 연구방법_history"].append(content)
            save_section_content("6. 연구방법", edited_content)
            st.success("편집된 내용이 저장되었습니다.")
            st.rerun()
    with col2:
        if st.button("실행 취소", key="undo_edit_6"):
            if st.session_state["6. 연구방법_history"]:
                # 히스토리에서 마지막 항목을 가져와 현재 내용으로 설정
                previous_content = st.session_state["6. 연구방법_history"].pop()
                save_section_content("6. 연구방법", previous_content)
                st.success("이전 버전으로 되돌렸습니다.")
                st.rerun()
            else:
                st.warning("더 이상 되돌릴 수 있는 버전이 없습니다.")

    # 글자 수 표시
    if content:
        char_count = len(content)
        st.info(f"현재 글자 수: {char_count}/500")
        if char_count > 500:
            st.warning("글자 수가 500자를 초과했습니다. 내용을 줄여주세요.")

# 7. 연구 과제명 작성 함수
def write_research_title():
    st.markdown("## 7. 연구 과제명")
    
    # 히스토리 초기화
    if "7. 연구 과제명_history" not in st.session_state:
        st.session_state["7. 연구 과제명_history"] = []

    # 완료 메시지 표시 (있는 경우)
    if 'completion_message' in st.session_state:
        st.success(st.session_state.completion_message)
        st.info("사이드바의 '전체 내용 미리보기' 버튼을 눌러 전체 내용을 확인하고 클립보드에 복사할 수 있습니다.")
        del st.session_state.completion_message

    # 안내 글 추가
    st.markdown("""
    연구 과제명을 직접 입력하거나, AI에게 추천받을 수 있습니다. 
                
    AI 추천은 기본적으로 3쌍의 영문/한글 제목을 제시합니다.
    
    AI 추천을 받으려면 '연구 과제명 추천받기' 버튼을 클릭하세요.
    """)

    # 사용자 입력 받기
    user_input = st.text_area("""연구 과제명에 대해 AI에게 알려줄 추가 정보나 고려사항이 있다면 입력해주세요. \n없다면 빈칸으로 두어도 됩니다. \n빈칸이라면 자동으로 이전 섹션들의 내용을 종합하여 알맞은 제목을 추천합니다.:""", height=150)

    # "연구 과제명 추천받기" 버튼을 여기로 이동
    if st.button("연구 과제명 AI에게 추천받기"):
        research_purpose = load_section_content("1. 연구 목적")
        research_background = load_section_content("2. 연구 배경")
        selection_criteria = load_section_content("3. 선정기준, 제외기준")
        sample_size = load_section_content("4. 대상자 수 및 산출근거")
        data_analysis = load_section_content("5. 자료분석과 통계적 방법")
        research_method = load_section_content("6. 연구방법")
        
        prompt = PREDEFINED_PROMPTS["7. 연구 과제명"].format(
            user_input=user_input,
            research_purpose=research_purpose,
            research_background=research_background,
            selection_criteria=selection_criteria,
            sample_size=sample_size,
            data_analysis=data_analysis,
            research_method=research_method
        )
        
        ai_response = generate_ai_response(prompt)
        
        # AI 응답 파싱 및 검증
        options = parse_and_validate_titles(ai_response)
        
        if options:
            save_section_content("7. 연구 과제명", "\n\n".join(options))
            st.rerun()
        else:
            st.error("AI가 올바른 형식의 연구 과제명을 생성하지 못했습니다. 다시 시도해주세요.")

    content = load_section_content("7. 연구 과제명")

    # 편집 기능 (항상 표시)
    edited_content = st.text_area(
        "연구 과제명을 직접 입력하거나 편집하세요:",
        content if content else "",
        height=150,
        key="edit_content_7"
    )

    col1, col2 = st.columns(2)
    with col1:
        if st.button("편집 내용 저장", key="save_edit_7"):
            if edited_content:
                # 현재 내용을 히스토리에 추가
                if content:
                    st.session_state["7. 연구 과제명_history"].append(content)
                save_section_content("7. 연구 과제명", edited_content)
                st.success("편집된 내용이 저장되었습니다.")
                st.rerun()
            else:
                st.warning("저장할 내용을 입력해주세요.")

    with col2:
        if st.button("이전 버전으로 되돌리기", key="undo_edit_7"):
            if st.session_state["7. 연구 과제명_history"]:
                previous_content = st.session_state["7. 연구 과제명_history"].pop()
                save_section_content("7. 연구 과제명", previous_content)
                st.success("이전 버전으로 되돌렸습니다.")
                st.rerun()
            else:
                st.warning("더 이상 되돌릴 수 있는 버전이 없습니다.")

    if content:
        options = content.split("\n\n")
        valid_options = [opt for opt in options if is_valid_title_option(opt)]

        if valid_options:
            st.markdown("### AI가 추천한 연구 과제명 (선택해주세요):")
            selected_option = st.radio(
                "",
                valid_options,
                format_func=lambda x: x.replace('\n', ' / '),  # 라디오 버튼에서는 간단히 표시
                index=0
            )
            st.markdown(format_title_option(selected_option), unsafe_allow_html=True)
            
            if st.button("선택한 연구 과제명 저장"):
                save_section_content("1. 연구 과제명", selected_option)
                st.session_state.completion_message = "IRB 연구 계획서 작성이 완료되었습니다!"
                st.rerun()
        else:
            st.error("유효한 연구 과제명 옵션이 없습니다. '연구 과제명 추천받기' 버튼을 다시 클릭해주세요.")

    # 수정 요청 기능
    if st.button("수정 요청하기", key="request_modification_7"):
        st.session_state.show_modification_request_7 = True
        st.rerun()

    if st.session_state.get('show_modification_request_7', False):
        modification_request = st.text_area(
            "수정을 원하는 부분과 수정 방향을 설명해주세요:",
            height=150,
            key="modification_request_7"
        )
        if st.button("수정 요청 제출", key="submit_modification_7"):
            if modification_request:
                current_content = load_section_content("7. 연구 과제명")
                # 현재 내용을 히스토리에 추가
                st.session_state["7. 연구 과제명_history"].append(current_content)
                
                prompt = f"""
                현재 연구 과제명 옵션들:
                {current_content}

                사용자의 수정 요청:
                {modification_request}

                위의 수정 요청을 반영하여 연구 과제명을 수정해주세요. 다음 지침을 따라주세요:
                1. 영문 제목과 한글 제목을 각각 작성해주세요.
                2. 총 3가지의 제목 옵션을 제시해주세요.
                3. 각 옵션은 영문 제목과 한글 제목이 한 쌍을 이루어야 합니다.
                4. 제목은 연구의 핵심 내용을 간결하고 명확하게 표현해야 합니다.
                5. 제목은 연구의 목적, 대상, 방법 등을 포함할 수 있습니다.
                6. 영문 제목은 첫 글자만 대문자로 작성하세요. (예: Effect of...)
                7. 수정 요청을 최대한 반영하되, 전체적인 일관성을 유지하세요.
                
                수정된 3가지 연구 과제명 옵션을 작성해주세요. 각 옵션은 다음과 같은 형식으로 작성해주세요:
                [영문 제목]
                [한글 제목]

                [영문 제목]
                [한글 제목]

                [영문 제목]
                [한글 제목]
                """
                modified_response = generate_ai_response(prompt)
                
                save_section_content("7. 연구 과제명", modified_response)
                st.session_state.show_modification_request_7 = False
                st.rerun()
            else:
                st.warning("수정 요청 내용을 입력해주세요.")

def display_references():
    st.markdown("### 참고문헌")
    references = format_references(
        st.session_state.get('scholar_results', []),
        st.session_state.get('pdf_files', [])
    )
    for i, ref in enumerate(references, 1):
        st.markdown(f"{i}. {ref}")


def extract_pdf_metadata(pdf_file):
    try:
        text = extract_text_from_pdf(pdf_file)
        
        # 텍스트의 처음 부분만 사용 (API 토큰 제한을 고려)
        text_sample = text[:5000]
        
        prompt = f"""
        다음은 학술 논문의 일부입니다. 이 논문의 제목, 저자들(최대 3명까지), 저자들의 소속 기관(특히 한국 소속 여부), 출판 연도를 추출해주세요.
        결과는 다음 형식으로 작성해주세요:
        제목: [논문 제목]
        저자: [저자1], [저자2], [저자3]
        소속: [소속1], [소속2], [소속3] (한국 소속이 있다면 'Korean' 태그를 추가, Seoul도 한국으로 간주)
        연도: [출판 연도]
        
        정보를 찾을 수 없는 경우 다음과 같이 구체적으로 표시해주세요:
        - 제목을 찾을 수 없는 경우: "Unknown title"
        - 저자를 찾을 수 없는 경우: "Unknown authors"
        - 소속을 찾을 수 없는 경우: "Unknown affiliations"
        - 연도를 찾을 수 없는 경우: "Unknown year"
        
        논문 내용:
        {text_sample}
        """
        
        response = st.session_state.anthropic_client.messages.create(
            model="claude-3-5-sonnet-20241022",
            max_tokens=300,
            messages=[{"role": "user", "content": prompt}]
        )
        
        result = response.content[0].text
        
       # 결과 파싱
        title = re.search(r'제목: (.+)', result)
        authors = re.search(r'저자: (.+)', result)
        affiliations = re.search(r'소속: (.+)', result)
        year = re.search(r'연도: (.+)', result)
        is_korean = re.search(r'한국 소속 여부: (.+)', result)
        
        return {
            'title': title.group(1) if title else "Unknown title",
            'authors': authors.group(1) if authors else "Unknown authors",
            'affiliations': affiliations.group(1) if affiliations else "Unknown affiliations",
            'year': year.group(1) if year else "Unknown year",
            'is_korean': is_korean.group(1).lower() == '예' if is_korean else False
        }
    except Exception as e:
        print(f"Error extracting metadata from {pdf_file.name}: {str(e)}")
        return {
            'title': "Unknown title",
            'authors': "Unknown authors",
            'affiliations': "Unknown affiliations",
            'year': "Unknown year",
            'is_korean': False
        }
        
def confirm_metadata(extracted_info):
    st.write("추출된 메타데이터:")
    title = st.text_input("제목", value=extracted_info['title'])
    authors = st.text_input("저자", value=extracted_info['authors'])
    year = st.text_input("년도", value=extracted_info['year'])
    return f"{authors}. {title}. {year}."

def format_references(pdf_files):
    references = []
    for i, pdf_file in enumerate(pdf_files, start=1):
        metadata = extract_pdf_metadata(pdf_file)
        reference = f"{i}. {metadata['authors']}. {metadata['title']}. {metadata['year']}."
        references.append(reference)
    return references

def parse_and_validate_titles(response):
    lines = response.split('\n')
    options = []
    current_option = []
    
    for line in lines:
        if line.strip().startswith(('1.', '2.', '3.')):
            if current_option:
                options.append('\n'.join(current_option))
                current_option = []
        current_option.append(line.strip())
    
    if current_option:
        options.append('\n'.join(current_option))
    
    return [opt for opt in options if is_valid_title_option(opt)]

def is_valid_title_option(option):
    lines = option.split('\n')
    return len(lines) >= 2 and lines[0].strip() and lines[1].strip()

def format_title_option(option):
    lines = option.split('\n')
    if len(lines) >= 2:
        return f"<p><strong>영문:</strong> {lines[0]}<br><strong>한글:</strong> {lines[1]}</p>"
    return f"<p>{option}</p>"
            

def extract_references(text):
    # [저자, 연도] 형식의 참고문헌을 추출
    references = re.findall(r'\[([^\]]+)\]', text)
    # 각 참고문헌을 [저자, 연도] 형식의 리스트로 변환
    return [ref.split(',') for ref in set(references)]

# 전체 인터페이스
def chat_interface():
    st.subheader("IRB 연구계획서 작성 도우미✏️ ver.02 (by HJY)")

    if 'current_research_id' not in st.session_state:
        st.session_state.current_research_id = generate_research_id()

    if 'view_mode' not in st.session_state:
        st.session_state.view_mode = 'edit'

    # 버튼
    if 'api_key' not in st.session_state or not st.session_state.api_key:
        api_key = st.text_input("Anthropic API 키를 입력하세요:", type="password")
        
        # API 키 확인 버튼
        if st.button("API 키 확인"):
            client = initialize_anthropic_client(api_key)
            if client:
                st.success("유효한 API 키입니다. 연구계획서 작성하기 버튼을 눌러 시작하세요.")
                st.session_state.temp_api_key = api_key
            else:
                st.error("API 키 설정에 실패했습니다. 키를 다시 확인해 주세요.")

        st.write("")  # 한 줄의 간격만 유지
        
        # 연구계획서 작성하기 버튼
        if st.button("연구계획서 작성하기 ✏️"):
            if 'temp_api_key' in st.session_state:
                st.session_state.api_key = st.session_state.temp_api_key
                st.session_state.anthropic_client = initialize_anthropic_client(st.session_state.api_key)
                del st.session_state.temp_api_key
                st.success("API 키가 설정되었습니다!")
                st.rerun()
            else:
                st.warning("먼저 API 키를 입력하고 확인해주세요.")

    # API 키가 설정된 후의 메인 인터페이스
    else:
        st.sidebar.text(f"현재 API 키: {st.session_state.api_key[:5]}...")

        if st.sidebar.button("🔄 초기화면으로"):
            for key in list(st.session_state.keys()):
                del st.session_state[key]
            st.rerun()

        if st.sidebar.button("새 연구계획서 시작"):
            reset_session_state()
            st.success("새로운 연구계획서를 시작합니다.")
            st.rerun()

        # 전체 내용 미리보기 버튼 추가
        if st.sidebar.button("전체 내용 미리보기"):
            with st.spinner('미리보기 모드로 전환 중...'):
                st.session_state.view_mode = 'preview'
                st.rerun()

        if 'current_section' not in st.session_state:
            st.session_state.current_section = 'home'

        # 조건부 렌더링
        if st.session_state.view_mode == 'edit':
            render_edit_mode()
        else:
            render_preview_mode()

def render_edit_mode():
    if st.session_state.current_section == 'home':
        render_home_page()
    else:
        render_section_page()

def render_home_page():
    st.markdown("## 연구계획서 작성을 시작합니다")
    st.markdown("아래 버튼을 클릭하여 각 섹션을 작성하세요. 각 파트만 선택해서 작성도 가능하지만, 최상의 결과를 위해서는 연구 목적 세션부터 시작하여 어플이 제공하는 순서대로 작성하는 것을 가장 추천합니다.")
    
    for section in RESEARCH_SECTIONS:
        if st.button(f"{section} 작성하기"):
            st.session_state.current_section = section
            st.rerun()

def render_section_page():
    # 현재 섹션에 따른 작성 인터페이스 표시
    if st.session_state.current_section == "7. 연구 과제명":
        write_research_title()
    elif st.session_state.current_section == "1. 연구 목적":
        write_research_purpose()
    elif st.session_state.current_section == "2. 연구 배경":
        write_research_background()
    elif st.session_state.current_section == "3. 선정기준, 제외기준":
        write_selection_criteria()
    elif st.session_state.current_section == "4. 대상자 수 및 산출근거":
        write_sample_size()
    elif st.session_state.current_section == "5. 자료분석과 통계적 방법":
        write_data_analysis()
    elif st.session_state.current_section == "6. 연구방법":
        write_research_method()

    # 이전 섹션과 다음 섹션 버튼
    col1, col2 = st.columns(2)
    
    with col1:
        if st.button("⬅️이전 섹션"):
            current_index = RESEARCH_SECTIONS.index(st.session_state.current_section)
            if current_index > 0:
                st.session_state.current_section = RESEARCH_SECTIONS[current_index - 1]
            else:
                st.session_state.current_section = 'home'
            st.rerun()

    with col2:
        if st.session_state.current_section != RESEARCH_SECTIONS[-1]:
            if st.button("다음 섹션➡️"):
                current_index = RESEARCH_SECTIONS.index(st.session_state.current_section)
                if current_index < len(RESEARCH_SECTIONS) - 1:
                    st.session_state.current_section = RESEARCH_SECTIONS[current_index + 1]
                    st.rerun()

    # 홈으로 돌아가기 버튼
    if st.button("홈으로 돌아가기"):
        st.session_state.current_section = 'home'
        st.rerun()

def render_preview_mode():
    st.markdown("## 전체 연구계획서 미리보기")
    
    sections_content = generate_full_content()

    for section, content in sections_content.items():
        st.subheader(section)
        if section == "참고문헌":
            references_content = st.text_area("참고문헌 편집", content, height=300)
            if st.button("참고문헌 저장", key="save_references"):
                save_section_content("참고문헌", references_content)
                st.success("참고문헌 내용이 저장되었습니다.")
        else:
            st.code(content, language="markdown")
    
    uploaded_file = st.file_uploader("IRB 연구계획서 DOCX 템플릿을 업로드하세요", type="docx")
    
    if uploaded_file is not None:
        if 'doc' not in st.session_state:
            st.session_state.doc = Document(uploaded_file)
        
        if st.button("업로드한 파일의 섹션 확인하기"):
            st.session_state.matching_results = {}
            for section in sections_content.keys():
                match = find_best_match(st.session_state.doc, section)
                st.session_state.matching_results[section] = match.text if match else "Not found"


            st.subheader("섹션 매칭 결과")
            for section, match_text in st.session_state.matching_results.items():
                st.write(f"{section}: {match_text}")
        
            st.session_state.show_confirm_button = True

        if st.session_state.get('show_confirm_button', False):
            if st.button("DOCX 파일 생성"):
                st.text("DOCX 파일 생성 버튼이 클릭되었습니다.")  # 버튼 클릭 확인 메시지
                try:
                    # 원본 템플릿의 복사본을 만듭니다
                    filled_doc = Document(BytesIO(uploaded_file.getvalue()))
                    st.text("DOCX 템플릿이 로드되었습니다.")
                    
                    # 복사본에 내용을 채웁니다
                    filled_doc = fill_docx_template(filled_doc, sections_content)
                    st.text("fill_docx_template 함수가 실행되었습니다.")
                    
                    # 채워진 문서를 메모리에 저장합니다
                    docx_file = BytesIO()
                    filled_doc.save(docx_file)
                    docx_file.seek(0)
                    
                    st.download_button(
                        label="완성된 DOCX 파일 다운로드",
                        data=docx_file,
                        file_name="완성된_연구계획서.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                    )
                    st.success("DOCX 파일이 성공적으로 생성되었습니다.")
                except Exception as e:
                    st.error(f"DOCX 파일 생성 중 오류가 발생했습니다: {str(e)}")
                    st.error("자세한 오류 정보:")
                    st.exception(e)

    if st.button("편집 모드로 돌아가기"):
        st.session_state.view_mode = 'edit'
        st.rerun()

def generate_full_content():
    sections_content = {}
        
    # 섹션 순서대로 내용 추가
    for section in RESEARCH_SECTIONS:
        section_content = load_section_content(section)
        if section_content:
            sections_content[section] = section_content
    
    # 참고문헌 편집된 내용을 세션 상태에서 로드하여 반영
    references_content = load_section_content("참고문헌")
    if references_content:
        sections_content["참고문헌"] = references_content
    else:
        # 기존 PDF 파일에서 참고문헌을 생성하는 로직 유지
        references = format_references(st.session_state.get('pdf_files', []))
        sections_content["참고문헌"] = "\n".join(references)
        # 참고문헌 세션 상태에 저장
        save_section_content("참고문헌", sections_content["참고문헌"])
    
    return sections_content

# IRB 템플릿 docx 파일 업로드
def upload_docx_template():
    uploaded_file = st.file_uploader("IRB 연구계획서 DOCX 템플릿을 업로드하세요", type="docx")
    if uploaded_file is not None:
        return Document(uploaded_file)
    return None

def normalize_text(text):
    # 숫자, 대소문자 구분 제거, 공백 및 특수 문자 제거
    return re.sub(r'\W+|\d+', '', text.lower())

def similarity_score(a, b):
    return SequenceMatcher(None, a, b).ratio()

def find_best_match(doc, section_title):
    normalized_section = normalize_text(section_title)
    best_match = None
    best_score = 0

    for paragraph in doc.paragraphs:
        normalized_para = normalize_text(paragraph.text)
        
        # 완전 일치 시 즉시 반환
        if normalized_section == normalized_para:
            return paragraph
        
        # 부분 일치 확인
        if normalized_section in normalized_para:
            score = similarity_score(normalized_section, normalized_para)
            if score > best_score:
                best_score = score
                best_match = paragraph

    # 유사도 임계값 (예: 0.7)
    if best_score > 0.7:
        return best_match
    
    return None

def insert_content_after_section(doc, section_title, content):
    section_para = find_best_match(doc, section_title)
    if section_para:
        # 현재 단락의 인덱스를 찾습니다
        index = doc.paragraphs.index(section_para)
        # 새 단락을 다음 인덱스에 추가합니다
        new_para = doc.add_paragraph()
        doc.paragraphs.insert(index + 1, new_para)
        new_para.text = content
        return new_para
    return None

def fill_docx_template(doc, sections_content):
    st.text("fill_docx_template 함수가 호출되었습니다.")  # 함수 호출 확인
    for section, content in sections_content.items():
        try:
            section_para = find_best_match(doc, section)  # 섹션 제목 단락 찾기
            if section_para:
                st.text(f"섹션 '{section}'을(를) 찾았습니다.")
                
                # 섹션 제목 아래에 새로운 내용 추가
                st.text(f"섹션 '{section}' 아래에 내용 추가 시도 중...")
                
                # 기존 섹션 제목 단락에 새로운 Run 추가
                new_run = section_para.add_run("\n" + content)
                
                # 스타일 설정 시 오류 방지: Run 객체에는 Paragraph 스타일을 사용하지 않음
                # new_run.style = 'Normal'  # 이 부분을 제거
                
                st.text(f"'{section}' 섹션 아래에 내용이 성공적으로 추가되었습니다.")
            else:
                st.warning(f"섹션 '{section}'을(를) 템플릿에서 찾을 수 없습니다.")
        except Exception as e:
            st.error(f"섹션 '{section}'에 내용을 추가하는 중 오류가 발생했습니다: {str(e)}")
    return doc

def download_docx(doc):
    # 메모리 상의 파일 객체 생성
    docx_file = BytesIO()
    # 문서를 메모리 상의 파일 객체에 저장
    doc.save(docx_file)
    # 파일 포인터를 처음으로 이동
    docx_file.seek(0)
    return docx_file

    # CSS 스타일
    st.markdown("""
    <style>
    .stButton>button {
        width: 100%;
        height: 60px;
        white-space: normal;
        word-wrap: break-word;
        text-align: center;
    }
    </style>
    """, unsafe_allow_html=True)

# 메인 함수 호출
if __name__ == "__main__":
    chat_interface()